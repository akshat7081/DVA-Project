"""
Generate Word document (.docx) matching Daksh's exact format, then convert to PDF.
Same fonts, same sizes, same spacing, same layout — pixel-perfect match.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start','top','end','bottom','insideH','insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr in kwargs[edge]:
                element.set(qn(f'w:{attr}'), str(kwargs[edge][attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_header(doc):
    """Add header: DVA Assignment 2 | Akshat Tripathi | 04914202023"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    
    # Create a table in header for 3-column layout
    tbl = header.add_table(rows=1, cols=3, width=Cm(17))
    tbl.autofit = True
    
    # Left: DVA Assignment 2
    cell_l = tbl.cell(0, 0)
    p_l = cell_l.paragraphs[0]
    run_l = p_l.add_run("DVA Assignment 2")
    run_l.font.name = "Times New Roman"
    run_l.font.size = Pt(10)
    run_l.bold = True
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Center: Akshat Tripathi
    cell_c = tbl.cell(0, 1)
    p_c = cell_c.paragraphs[0]
    run_c = p_c.add_run("Akshat Tripathi")
    run_c.font.name = "Calibri"
    run_c.font.size = Pt(11)
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Right: 04914202023
    cell_r = tbl.cell(0, 2)
    p_r = cell_r.paragraphs[0]
    run_r = p_r.add_run("04914202023")
    run_r.font.name = "Calibri"
    run_r.font.size = Pt(11)
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Remove borders from header table
    for cell in tbl.row_cells(0):
        set_cell_border(cell,
            top={"sz":"0","val":"none","color":"auto"},
            bottom={"sz":"0","val":"none","color":"auto"},
            start={"sz":"0","val":"none","color":"auto"},
            end={"sz":"0","val":"none","color":"auto"})

def add_para(doc, text, font_name="Times New Roman", font_size=12, bold=False, 
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    
    # Handle bold sections within text using <b> tags
    if '<b>' in text:
        parts = text.split('<b>')
        for i, part in enumerate(parts):
            if '</b>' in part:
                bold_text, normal_text = part.split('</b>', 1)
                run_b = p.add_run(bold_text)
                run_b.font.name = font_name
                run_b.font.size = Pt(font_size)
                run_b.bold = True
                if normal_text:
                    run_n = p.add_run(normal_text)
                    run_n.font.name = font_name
                    run_n.font.size = Pt(font_size)
            else:
                run = p.add_run(part)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = bold
    else:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
    return p

def add_bullet(doc, text, font_name="Times New Roman", font_size=12, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    
    if indent_level == 0:
        # Use bullet symbol
        prefix = "\u2022 "  # bullet •
        p.paragraph_format.left_indent = Cm(1.0)
    else:
        prefix = "o "
        p.paragraph_format.left_indent = Cm(2.0)
    
    run = p.add_run(prefix + text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return p

def add_code_block(doc, code_text, font_name="Courier New", font_size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(code_text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return p

def build_docx():
    doc = Document()
    
    # Set margins matching Daksh's
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    add_header(doc)
    
    # ════════════════════════════════════════
    # PAGE 1 — Cover + Project Overview
    # ════════════════════════════════════════
    add_para(doc, "", space_after=Pt(40))  # spacer
    add_para(doc, "Assignment-2", "Times New Roman", 24, True, WD_ALIGN_PARAGRAPH.CENTER, Pt(12))
    add_para(doc, "", space_after=Pt(8))
    add_para(doc, "Data Visualization and Analytics Project: Social Media Trends Analytics Dashboard",
             "Times New Roman", 14, True, WD_ALIGN_PARAGRAPH.CENTER, Pt(12))
    add_para(doc, "", space_after=Pt(8))
    
    add_para(doc, "Project Link: https://github.com/akshattripathi1/Social-Media-Trends-Dashboard",
             "Times New Roman", 12, False, space_after=Pt(4))
    add_para(doc, "Dataset: Custom Social Media Trends Dataset (2000 records, 12 columns)",
             "Times New Roman", 12, False, space_after=Pt(12))
    
    add_para(doc, "Project Overview", "Times New Roman", 12, True, space_after=Pt(4), space_before=Pt(8))
    add_para(doc, 
        "This project presents a <b>multi-page interactive dashboard</b> developed using Python, Tkinter, Pandas, and "
        "Matplotlib to analyze <b>social media trends and engagement metrics</b> across different platforms.",
        space_after=Pt(4))
    add_para(doc,
        "The system transforms raw social media data into a <b>visual analytics platform</b> that helps users understand "
        "platform performance, content engagement patterns, and trending hashtags.",
        space_after=Pt(4))
    add_para(doc, "The dashboard is designed to provide insights into:", space_after=Pt(2))
    
    for b in [
        "Social media engagement trends across Twitter, Instagram, YouTube, Facebook, and TikTok",
        "Hashtag performance and trending analysis across 10 content categories",
        "Platform-wise reach comparison (Views, Likes, Comments, Shares)",
        "Content category engagement patterns (Entertainment, Sports, Technology, etc.)",
        "Metric correlation analysis (Posts vs Likes vs Views vs Engagement Rate)",
        "Platform \u00d7 Category engagement heatmap for identifying content sweet spots",
    ]:
        add_bullet(doc, b)
    
    add_para(doc,
        "The interface follows a <b>clean modern UI design</b> with a midnight purple theme and color-coded "
        "platform indicators for better visual understanding.",
        space_after=Pt(8))
    
    # Dataset Overview
    add_para(doc, "Dataset Overview", "Times New Roman", 12, True, space_after=Pt(4), space_before=Pt(8))
    add_para(doc,
        "The dataset contains <b>2000 records of social media trend data</b> simulating realistic engagement "
        "patterns across multiple platforms and content categories.",
        space_after=Pt(4))
    add_para(doc, "It includes:", space_after=Pt(2))
    
    for b in [
        "Platform (Twitter, Instagram, YouTube, Facebook, TikTok)",
        "Hashtag name (60 unique hashtags)",
        "Content Category (Entertainment, Sports, Politics, Technology, Fashion, Gaming, Food, Travel, Education, Health)",
        "Country of origin (11 countries, India-biased)",
        "Date and Month",
        "Metrics:",
    ]:
        add_bullet(doc, b)
    for sb in ["Posts count", "Likes", "Views", "Comments", "Shares", "Engagement Rate (%)"]:
        add_bullet(doc, sb, indent_level=1)
    
    # ════════════════════════════════════════
    # PAGE 2 — Tech Stack (continues naturally)
    # ════════════════════════════════════════
    add_para(doc, "This dataset is used to analyze social media trends and visualize engagement patterns effectively.",
             space_after=Pt(8))
    
    add_para(doc, "Tech Stack", "Times New Roman", 12, True, space_after=Pt(4), space_before=Pt(8))
    for tech, desc in [
        ("Python", "Programming language"),
        ("Tkinter", "GUI framework for dashboard"),
        ("Pandas", "Data cleaning and analysis"),
        ("Matplotlib", "Data visualization (charts & graphs)"),
        ("NumPy", "Numerical computation and data processing"),
    ]:
        add_para(doc, f"<b>{tech}</b> \u2192 {desc}", space_after=Pt(3))
    
    doc.add_page_break()
    
    # ════════════════════════════════════════
    # PAGES 3-18 — Program Code
    # ════════════════════════════════════════
    add_para(doc, "Program Code", "Times New Roman", 12, True, space_after=Pt(8), space_before=Pt(4))
    
    with open("dashboard.py", "r", encoding="utf-8") as f:
        code_lines = f.read().split("\n")
    
    # Add code in chunks — Daksh had ~30 lines per page
    chunk_size = 16
    for i in range(0, len(code_lines), chunk_size):
        chunk = "\n".join(code_lines[i:i+chunk_size])
        add_code_block(doc, chunk, "Courier New", 8)
        if i + chunk_size < len(code_lines):
            doc.add_page_break()
    
    doc.add_page_break()
    
    # ════════════════════════════════════════
    # PAGES 19-21 — Output Screenshots
    # ════════════════════════════════════════
    add_para(doc, "Output", "Times New Roman", 16, True, space_after=Pt(8))
    
    ss_pages = [
        [("page1_overview.png", "Page 1- Overview"),
         ("page2_platform.png", "Page 2- Platform Analysis")],
        [("page3_category.png", "Page 3- Category Trends"),
         ("page4_hashtag.png", "Page 4- Hashtag Rankings")],
        [("page5_heatmap.png", "Page 5- Heatmap & Correlations")],
    ]
    
    for pg_idx, page_group in enumerate(ss_pages):
        for fname, caption in page_group:
            fpath = os.path.join("screenshots", fname)
            if os.path.exists(fpath):
                add_para(doc, caption, "Times New Roman", 12, True, space_after=Pt(4))
                doc.add_picture(fpath, width=Cm(16))
                add_para(doc, "", space_after=Pt(8))
        if pg_idx < len(ss_pages) - 1:
            doc.add_page_break()
    
    # Save DOCX
    docx_path = "DVA_Assignment_2_Akshat_Tripathi.docx"
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")
    
    # Convert to PDF
    try:
        from docx2pdf import convert
        convert(docx_path, "DVA_Assignment_2_Akshat_Tripathi.pdf")
        print("PDF generated from Word!")
    except Exception as e:
        print(f"Auto-convert failed ({e}). Open the .docx in Word and Save As PDF manually.")

if __name__ == "__main__":
    build_docx()
